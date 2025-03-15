from docx import Document
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt  # Исправленный импорт Pt
import os

# Путь к папке с вашими Python-файлами
folder_path = "C:/Users/User/PycharmProjects/telegram_service/app"  # Указанная вами папка
output_file = "all_python_files.docx"  # Имя выходного файла

# Создаём новый документ Word
doc = Document()

# Добавляем стиль "Code", если его нет
styles = doc.styles
if "Code" not in styles:
    style = styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = "Courier New"  # Моноширинный шрифт для кода
    style.font.size = Pt(10)  # Размер шрифта 10pt
    style.paragraph_format.space_after = Pt(6)  # Отступ после абзаца 6pt

# Добавляем заголовок документа
doc.add_heading("Список Python-файлов", level=1)

# Проходим по всем файлам в папке и её подпапках
for root, dirs, files in os.walk(folder_path):
    # Пропускаем папку __pycache__
    if "__pycache__" in dirs:
        dirs.remove("__pycache__")  # Удаляем __pycache__ из списка обрабатываемых директорий

    for file in files:
        if file.endswith(".py"):  # Обрабатываем только .py файлы
            file_path = os.path.join(root, file)
            doc.add_heading(f"Файл: {file_path}", level=2)

            # Читаем содержимое файла
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    content = f.read()
                    doc.add_paragraph(content, style="Code")  # Добавляем код с форматированием
            except Exception as e:
                doc.add_paragraph(f"Ошибка чтения файла {file_path}: {e}", style="Normal")

            doc.add_page_break()  # Разделяем файлы новой страницей

# Сохраняем документ
doc.save(output_file)
print(f"Файлы сохранены в {output_file}")